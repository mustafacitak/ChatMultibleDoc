import os
import uuid
import yaml
import functools
import sys
import json
import requests
import fitz  # PyMuPDF
import docx
import numpy as np
from bs4 import BeautifulSoup
from dotenv import load_dotenv
import streamlit as st
from streamlit import dialog
import PyPDF2
import pandas as pd
from paddleocr import PaddleOCR
from PIL import Image
import io
import asyncio
from utils.net_search import net_search_async
from utils.hash_utils import load_chunks_from_jsonl, save_chunks_to_jsonl
# LangChain loaders - güncel import yolu
from langchain_community.document_loaders import UnstructuredExcelLoader, CSVLoader

# .env dosyasını yükle
load_dotenv()

# API anahtarını oku (öncelik sırası: GOOGLE_API_KEY > GEMINI_API_KEY)
GEMINI_API_KEY = os.getenv('GOOGLE_API_KEY') or os.getenv('GEMINI_API_KEY')

# Modül arama yolunu ayarla - dosyanın konumunu ana dizin olarak ekle
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from loguru import logger
logger.remove()
logger.add(sys.stderr, level="INFO")
logger.add("terminal_log.txt", rotation="5 MB", level="DEBUG")

from utils.hash_utils import calculate_chunk_hash, get_file_hash, is_file_already_processed, get_existing_chunk_hashes
from utils.collection_utils import normalize_collection_name, get_existing_collections, create_or_update_collection
from utils.chunking_utils import dynamic_chunking, sentence_chunking
from utils.embedding_utils import generate_embedding
from utils.duplicate_utils import is_duplicate_by_embedding, get_duplicate_threshold_from_config

# Hata yakalama dekoratörleri
def error_handler(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            logger.error(f"Hata oluştu - Fonksiyon: {func.__name__}, Hata: {e}")
            st.error(f"İşlem sırasında bir hata oluştu: {e}")
            return None
    return wrapper

def async_error_handler(func):
    @functools.wraps(func)
    async def wrapper(*args, **kwargs):
        try:
            return await func(*args, **kwargs)
        except Exception as e:
            logger.error(f"Async hata oluştu - Fonksiyon: {func.__name__}, Hata: {e}")
            st.error(f"İşlem sırasında bir hata oluştu: {e}")
            return None
    return wrapper

# --- Metin Çıkarma Fonksiyonları ---
@error_handler
def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF dosyasından metin çıkarılırken hata: {e}")
        st.error(f"PDF dosyasından metin çıkarılırken hata: {e}")
        return ""

@error_handler
def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"DOCX dosyasından metin çıkarılırken hata: {e}")
        st.error(f"DOCX dosyasından metin çıkarılırken hata: {e}")
        return ""

@error_handler
def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        logger.error(f"TXT dosyasından metin çıkarılırken hata: {e}")
        st.error(f"TXT dosyasından metin çıkarılırken hata: {e}")
        return ""

@error_handler
def extract_text_from_csv(file_path):
    try:
        # LangChain'in CSVLoader sınıfını kullan
        loader = CSVLoader(file_path=file_path)
        documents = loader.load()
        
        if not documents:
            logger.warning(f"CSVLoader ile dosyadan belge çıkarılamadı: {file_path}")
            return ""
        
        # İlk birkaç satırı içerik önizlemesi olarak al
        max_preview_rows = 20
        preview_text = "\n".join([doc.page_content for doc in documents[:max_preview_rows]])
        
        # Tüm belgelerin metinlerini birleştir
        full_text = "\n".join([doc.page_content for doc in documents])
        
        # Dosya çok büyükse bilgilendirici metin ekle
        if len(documents) > max_preview_rows:
            full_text += f"\n\n[... {len(documents)} satırlık CSV dosyasından tamamı işlendi ...]"
        
        logger.info(f"CSV başarıyla okundu: {len(documents)} satır içeriyor.")
        return full_text
    except Exception as e:
        logger.error(f"CSVLoader ile CSV dosyasından metin çıkarılırken hata: {e}")
        st.error(f"CSV dosyasından metin çıkarılırken hata: {e}")
        
        # Alternatif yöntem - Pandas ile deneme
        try:
            logger.info("Alternatif yöntem ile CSV dosyası okunmaya çalışılıyor (pandas)...")
            df = pd.read_csv(file_path, encoding='utf-8')
            text = df.to_string(index=False)
            logger.info(f"Pandas ile CSV başarıyla okundu: {len(df)} satır, {len(df.columns)} sütun")
            return text
        except Exception as e2:
            logger.error(f"Pandas ile CSV okuma hatası: {e2}")
            return ""

@error_handler
def extract_text_from_excel(file_path):
    try:
        # LangChain'in UnstructuredExcelLoader sınıfını kullan
        loader = UnstructuredExcelLoader(file_path=file_path)
        documents = loader.load()
        
        if not documents:
            logger.warning(f"UnstructuredExcelLoader ile dosyadan belge çıkarılamadı: {file_path}")
            return ""
        
        # İlk birkaç satırı içerik önizlemesi olarak al
        max_preview_rows = 20
        preview_text = "\n".join([doc.page_content for doc in documents[:max_preview_rows]])
        
        # Tüm belgelerin metinlerini birleştir
        full_text = "\n".join([doc.page_content for doc in documents])
        
        # Dosya çok büyükse bilgilendirici metin ekle
        if len(documents) > max_preview_rows:
            full_text += f"\n\n[... {len(documents)} satırlık Excel dosyasından tamamı işlendi ...]"
        
        logger.info(f"Excel başarıyla okundu: {len(documents)} satır içeriyor.")
        return full_text
    except Exception as e:
        logger.error(f"UnstructuredExcelLoader ile Excel dosyasından metin çıkarılırken hata: {e}")
        st.error(f"Excel dosyasından metin çıkarılırken hata: {e}")
        
        # Alternatif yöntem - Pandas ile deneme
        try:
            logger.info("Alternatif yöntem ile Excel dosyası okunmaya çalışılıyor (pandas)...")
            df = pd.read_excel(file_path)
            text = df.to_string(index=False)
            logger.info(f"Pandas ile Excel başarıyla okundu: {len(df)} satır, {len(df.columns)} sütun")
            return text
        except Exception as e2:
            logger.error(f"Pandas ile Excel okuma hatası: {e2}")
            return ""

@error_handler
def extract_text_from_url(url):
    try:
        response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})
        soup = BeautifulSoup(response.text, 'html.parser')
        for script in soup(["script", "style"]):
            script.extract()
        text = soup.get_text(separator="\n")
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text
    except Exception as e:
        logger.error(f"URL'den metin çıkarılırken hata: {e}")
        st.error(f"URL'den metin çıkarılırken hata: {e}")
        return ""

@error_handler
def extract_text_from_image(image_path):
    try:
        ocr = PaddleOCR(use_angle_cls=True, lang='tr')
        result = ocr.ocr(image_path, cls=True)
        texts = []
        for line in result:
            for item in line:
                texts.append(item[1][0])
        return "\n".join(texts)
    except Exception as e:
        logger.error(f"Görselden OCR ile metin çıkarılırken hata: {e}")
        st.error(f"Görselden OCR ile metin çıkarılırken hata: {e}")
        return ""

@error_handler
def extract_images_text_from_pdf_with_paddleocr(pdf_path):
    try:
        ocr = PaddleOCR(use_angle_cls=True, lang='tr')
        doc = fitz.open(pdf_path)
        text_from_images = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image = Image.open(io.BytesIO(image_bytes))
                img_np = np.array(image)
                result = ocr.ocr(img_np, cls=True)
                for line in result:
                    for item in line:
                        text_from_images.append(item[1][0])
        return "\n".join(text_from_images)
    except Exception as e:
        logger.error(f"PDF'deki görsellerden OCR ile metin çıkarılırken hata: {e}")
        st.error(f"PDF'deki görsellerden OCR ile metin çıkarılırken hata: {e}")
        return ""

@error_handler
def extract_images_text_from_docx_with_paddleocr(docx_path):
    try:
        ocr = PaddleOCR(use_angle_cls=True, lang='tr')
        doc = docx.Document(docx_path)
        text_from_images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image = Image.open(io.BytesIO(image_data))
                img_np = np.array(image)
                result = ocr.ocr(img_np, cls=True)
                for line in result:
                    for item in line:
                        text_from_images.append(item[1][0])
        return "\n".join(text_from_images)
    except Exception as e:
        logger.error(f"DOCX'deki görsellerden OCR ile metin çıkarılırken hata: {e}")
        st.error(f"DOCX'deki görsellerden OCR ile metin çıkarılırken hata: {e}")
        return ""

# Config
@error_handler
def load_app_config(config_path='config/chunking_config.yaml'):
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = yaml.safe_load(f)
        app_config = config.get('app', {})
        return {
            'default_collection': app_config.get('default_collection', "ornek_koleksiyon"),
            'collections_path': app_config.get('collections_path', "db/collections"),
            'model_name': None
        }
    except Exception as e:
        logger.error(f"App config yükleme hatası: {e}")
        return {
            'default_collection': "ornek_koleksiyon",
            'collections_path': "db/collections",
            'model_name': None
        }

app_config = load_app_config()
COLLECTION_PATH_BASE = app_config['collections_path']
DEFAULT_COLLECTION = app_config['default_collection']
CONFIG_PATH = None
DOCS_PATH = None
MEMORY_PATH = None
if "collection_name" not in st.session_state:
    st.session_state.collection_name = DEFAULT_COLLECTION

@error_handler
def load_or_create_config(config_path):
    default_config = {
        'chunking': 'sentence',
        'chunk_size': 5,
        'overlap': 0
    }
    if not os.path.exists(config_path):
        os.makedirs(os.path.dirname(config_path), exist_ok=True)
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(default_config, f, ensure_ascii=False, indent=2)
        logger.info(f"Varsayılan config.json oluşturuldu: {config_path}")
        return default_config
    else:
        with open(config_path, 'r', encoding='utf-8') as f:
            config = json.load(f)
        logger.info(f"Config yüklendi: {config_path}")
        return config

def process_uploaded_file(file_path, source_type):
    global CONFIG_PATH, DOCS_PATH, MEMORY_PATH
    CONFIG_PATH = os.path.join(COLLECTION_PATH_BASE, st.session_state.collection_name, "config.json")
    DOCS_PATH = os.path.join(COLLECTION_PATH_BASE, st.session_state.collection_name, "documents.jsonl")
    MEMORY_PATH = os.path.join(COLLECTION_PATH_BASE, st.session_state.collection_name, "memory.jsonl")
    config = load_or_create_config(CONFIG_PATH)
    file_hash = get_file_hash(file_path)
    if is_file_already_processed(file_hash, DOCS_PATH):
        logger.warning('Bu dosya zaten yüklenmiş, tekrar işlenmeyecek.')
        return
    if file_path.endswith('.pdf'):
        text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        text = extract_text_from_docx(file_path)
    elif file_path.endswith('.txt'):
        text = extract_text_from_txt(file_path)
    elif file_path.endswith('.csv'):
        text = extract_text_from_csv(file_path)
    elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
        text = extract_text_from_excel(file_path)
    else:
        logger.error('Desteklenmeyen dosya formatı!')
        return
    if config['chunking'] == 'sentence':
        chunks = sentence_chunking(text, chunk_size=config['chunk_size'])
    else:
        chunks = dynamic_chunking(text, chunk_size=config['chunk_size'], overlap=config['overlap'])
    chunk_dicts = []
    for i, chunk in enumerate(chunks):
        emb = generate_embedding(chunk)
        chunk_dicts.append({
            'id': f'chunk_{uuid.uuid4()}',
            'text': chunk,
            'embedding': emb,
            'source': os.path.basename(file_path),
            'source_type': source_type,
            'file_hash': file_hash
        })
    save_chunks_to_jsonl(chunk_dicts, DOCS_PATH)
    logger.info(f"{len(chunk_dicts)} chunk kaydedildi.")

def upload_document():
    st.title("📎 Upload or Search Document")
    existing_collections = get_existing_collections()
    if "show_collection_dialog" not in st.session_state:
        st.session_state.show_collection_dialog = False
    @st.dialog("Create New Collection")
    def create_collection_dialog():
        st.markdown("### Enter a name for the new collection")
        collection_name = st.text_input("Collection Name", key="new_collection_name", placeholder="e.g. finance_documents")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Create", key="create_button", type="primary"):
                if not collection_name.strip():
                    st.warning("Collection name cannot be empty!")
                else:
                    norm_name = normalize_collection_name(collection_name)
                    col_path = os.path.join(COLLECTION_PATH_BASE, norm_name)
                    os.makedirs(col_path, exist_ok=True)
                    config_path = os.path.join(col_path, "config.json")
                    if not os.path.exists(config_path):
                        config = {
                            'chunking': 'sentence',
                            'chunk_size': 5,
                            'overlap': 0
                        }
                        with open(config_path, 'w', encoding='utf-8') as f:
                            json.dump(config, f, ensure_ascii=False, indent=2)
                    st.session_state.collection_name = norm_name
                    st.session_state.show_collection_dialog = False
                    st.success(f"Collection '{norm_name}' created successfully!")
                    st.rerun()
        with col2:
            if st.button("❌ Cancel", key="cancel_button"):
                st.session_state.show_collection_dialog = False
                st.rerun()
    with st.sidebar:
        st.markdown("### Collection Management")
        if existing_collections:
            default_index = existing_collections.index(st.session_state.collection_name) if st.session_state.collection_name in existing_collections else 0
            selected_collection = st.selectbox(
                "Existing Collections",
                options=existing_collections,
                index=default_index,
                key="collection_selector"
            )
            st.session_state.collection_name = selected_collection
        else:
            st.warning("No collections yet. Please create a new collection.")
        if st.session_state.collection_name in existing_collections:
            st.success(f"Active collection: **{st.session_state.collection_name}**")
            col_path = os.path.join(COLLECTION_PATH_BASE, st.session_state.collection_name)
            doc_path = os.path.join(col_path, "documents.jsonl")
            if os.path.exists(doc_path):
                doc_count = 0
                try:
                    with open(doc_path, 'r', encoding='utf-8') as f:
                        for _ in f:
                            doc_count += 1
                    st.caption(f"📄 There are **{doc_count}** documents in the collection")
                except:
                    pass
        else:
            st.info(f"Active collection: **{st.session_state.collection_name}**")
            st.caption("❗ This collection has not been created yet. It will be created automatically when a document is added.")
        if st.button("✨ Create New Collection", key="new_collection_button", type="primary"):
            st.session_state.show_collection_dialog = True
            create_collection_dialog()
    tab1, tab2, tab3 = st.tabs(["📎 Upload File", "🔗🌐 Add URL/Net_Search", "🖼️ Image OCR"])
    def process_and_add_to_collection(text, source_name, source_type):
        global DOCS_PATH, CONFIG_PATH
        CONFIG_PATH = os.path.join(COLLECTION_PATH_BASE, st.session_state.collection_name, "config.json")
        DOCS_PATH = os.path.join(COLLECTION_PATH_BASE, st.session_state.collection_name, "documents.jsonl")
        config_path = CONFIG_PATH
        config = load_or_create_config(config_path)
        if config['chunking'] == 'sentence':
            chunks = sentence_chunking(text, chunk_size=config['chunk_size'])
        else:
            chunks = dynamic_chunking(text, chunk_size=config['chunk_size'], overlap=config['overlap'])
        existing_hashes = get_existing_chunk_hashes(st.session_state.collection_name)
        existing_embs = []
        duplicate_threshold = get_duplicate_threshold_from_config()
        if os.path.exists(DOCS_PATH):
            with open(DOCS_PATH, 'r', encoding='utf-8') as f:
                for line in f:
                    try:
                        obj = json.loads(line)
                        if 'embedding' in obj:
                            existing_embs.append(obj['embedding'])
                    except Exception:
                        continue
        new_embs = []
        chunk_dicts = []
        file_hash = str(uuid.uuid4())
        new_chunk_count = 0
        for i, chunk in enumerate(chunks):
            chunk_hash = calculate_chunk_hash(chunk)
            if chunk_hash in existing_hashes or calculate_chunk_hash(source_name) in existing_hashes:
                logger.info(f"Duplicate detected by hash and skipped.")
                continue
            emb = generate_embedding(chunk)
            if is_duplicate_by_embedding(emb, existing_embs + new_embs, threshold=duplicate_threshold):
                logger.info(f"Semantic duplicate detected by embedding and skipped.")
                continue
            chunk_dicts.append({
                'id': f'chunk_{uuid.uuid4()}',
                'text': chunk,
                'embedding': emb,
                'source': source_name,
                'source_type': source_type,
                'file_hash': file_hash,
                'chunk_hash': chunk_hash
            })
            new_embs.append(emb)
            new_chunk_count += 1
        if chunk_dicts:
            create_or_update_collection(chunk_dicts, st.session_state.collection_name, source_name=source_name)
        return new_chunk_count
    with tab1:
        uploaded_files = st.file_uploader("Select one or more files (PDF, DOCX, TXT, CSV, XLSX, XLS)", type=["pdf", "docx", "txt", "csv", "xlsx", "xls"], accept_multiple_files=True)
        if uploaded_files:
            if st.button("Upload and Process", key="upload_files_btn"):
                for uploaded_file in uploaded_files:
                    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
                    temp_path = f"temp_upload_{uuid.uuid4()}{file_ext}"
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    text = ""
                    ocr_text = ""
                    if file_ext == ".pdf":
                        text = extract_text_from_pdf(temp_path)
                        ocr_text = extract_images_text_from_pdf_with_paddleocr(temp_path)
                        source_type = "pdf"
                    elif file_ext == ".docx":
                        text = extract_text_from_docx(temp_path)
                        ocr_text = extract_images_text_from_docx_with_paddleocr(temp_path)
                        source_type = "docx"
                    elif file_ext == ".txt":
                        text = extract_text_from_txt(temp_path)
                        source_type = "txt"
                    elif file_ext == ".csv":
                        text = extract_text_from_csv(temp_path)
                        source_type = "csv"
                    elif file_ext in [".xlsx", ".xls"]:
                        text = extract_text_from_excel(temp_path)
                        source_type = "excel"
                    else:
                        st.error(f"{uploaded_file.name}: Unsupported file format!")
                        text = ""
                        source_type = ""
                    all_text = text
                    if ocr_text and ocr_text.strip():
                        all_text = (text + "\n" + ocr_text).strip()
                    if source_type and all_text:
                        with st.spinner(f"Processing {uploaded_file.name} ..."):
                            n = process_and_add_to_collection(all_text, uploaded_file.name, source_type)
                        st.success(f"{uploaded_file.name} successfully added to {st.session_state.collection_name} collection. New chunks added: {n}")
                    os.remove(temp_path)
    with tab2:
        st.markdown("**🔗🌐 Add URL/Net_Search**")
        input_col1, input_col2 = st.columns(2)
        with input_col1:
            url = st.text_input("Enter a web address:", key="url_input")
        with input_col2:
            web_query = st.text_input("Enter a query to search the web:", key="web_search_query")
        button_col1, button_col2 = st.columns(2)
        with button_col1:
            add_url = st.button("Fetch and Add from URL", key="url_btn")
        with button_col2:
            add_web = st.button("Search the Web and Add", key="web_search_btn")
        result_msg = None
        web_results = None
        combined_text = ""
        if add_url or add_web:
            texts = []
            sources = []
            if add_url and url:
                with st.spinner("Fetching content from URL ..."):
                    text = extract_text_from_url(url)
                    if text:
                        texts.append(text)
                        sources.append(url)
                    else:
                        result_msg = ("error", "Could not extract text from URL.")
            if add_web and web_query:
                with st.spinner("Searching the web ..."):
                    results = asyncio.run(net_search_async(web_query, max_results=5))
                    if results:
                        web_results = results
                        logger.info(f"Web search results: {len(results)} found")
                        web_text = "\n".join([f"{r['title']}\n{r['url']}" for r in results if r.get('title') and r.get('url')])
                        texts.append(web_text)
                        sources.append(f"web_search:{web_query}")
                    else:
                        result_msg = ("error", "No results from web search.")
            if texts:
                combined_text = "\n".join(texts)
                n = process_and_add_to_collection(combined_text, ", ".join(sources), "url_web_combined")
                result_msg = ("success", f"Data successfully added to {st.session_state.collection_name} collection. New chunks added: {n}")
        if result_msg:
            if result_msg[0] == "success":
                st.success(result_msg[1])
            else:
                st.error(result_msg[1])
        if web_results:
            with st.expander("Search Results"):
                for r in web_results:
                    st.markdown(f"- [{r.get('title','')}]({r.get('url','')})")
    with tab3:
        uploaded_imgs = st.file_uploader("Select one or more images or scanned PDFs (JPG, PNG, TIFF, BMP, PDF)", type=["jpg", "jpeg", "png", "tiff", "bmp", "pdf"], accept_multiple_files=True)
        if uploaded_imgs:
            if st.button("Process with OCR", key="ocr_btn"):
                for uploaded_img in uploaded_imgs:
                    file_ext = os.path.splitext(uploaded_img.name)[1].lower()
                    temp_path = f"temp_img_upload_{uuid.uuid4()}{file_ext}"
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_img.getbuffer())
                    if file_ext == ".pdf":
                        st.warning(f"{uploaded_img.name}: OCR from PDF is not yet supported. Please upload an image.")
                        text = ""
                    else:
                        text = extract_text_from_image(temp_path)
                    if text:
                        with st.spinner(f"Extracting and processing text from {uploaded_img.name} ..."):
                            n = process_and_add_to_collection(text, uploaded_img.name, "ocr")
                        st.success(f"{uploaded_img.name} successfully added to {st.session_state.collection_name} collection. New chunks added: {n}")
                    os.remove(temp_path)

upload_document() 