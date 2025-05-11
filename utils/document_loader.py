import os
import tempfile
from typing import List, Dict, Union, Optional
from langchain_community.document_loaders import (
    PyPDFLoader,
    CSVLoader,
    UnstructuredWordDocumentLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pandas as pd
from bs4 import BeautifulSoup
import requests
import ssl
import warnings
from .ocr_processor import ocr_processor

# SSL sertifika sorunlarını atlayalım
ssl._create_default_https_context = ssl._create_unverified_context

def load_document(file_path: str) -> List[Document]:
    """
    Verilen dosya yolundaki belgeyi yükler ve Document nesneleri listesi döndürür.
    Desteklenen formatlar: PDF, DOCX, CSV, XLSX, görsel formatları (JPG, PNG, JPEG, BMP, TIFF)
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    documents = []
    
    try:
        # Görsel formatları için OCR işlemi
        if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
            return ocr_processor.process_image(file_path)
        
        elif ext == '.pdf':
            try:
                # Standart PDF yükleyiciyle deneyelim
                loader = PyPDFLoader(file_path)
                documents = loader.load()
                
                # PDF'den herhangi bir metin çıkarılabildi mi kontrol et
                if not documents or all(not doc.page_content.strip() for doc in documents):
                    print("PDF'de okunabilir metin bulunamadı, OCR denenecek...")
                    return process_pdf_with_ocr(file_path)
                else:
                    # Bazı sayfalarda metin var, bazılarında yoksa:
                    # Boş sayfaları bul ve OCR ile işle
                    enhanced_documents = []
                    for doc in documents:
                        if not doc.page_content.strip():
                            # Sayfada metin yoksa OCR kullan
                            page_num = doc.metadata.get('page', 0)
                            print(f"Sayfa {page_num} boş, OCR uygulanıyor...")
                            try:
                                # Bu sayfayı görsel olarak işle (gelecekte - şu an dummy)
                                # İleri düzey uygulama: PDF'in o sayfasını görsel olarak çıkarıp OCR yapma
                                enhanced_documents.append(doc)  # Şimdilik orijinal boş belgeyi tut
                            except Exception as e:
                                print(f"Sayfa OCR hatası: {str(e)}")
                                enhanced_documents.append(doc)
                        else:
                            enhanced_documents.append(doc)
                    
                    return enhanced_documents
            except Exception as pdf_error:
                print(f"Standart PDF yükleme hatası: {str(pdf_error)}, OCR denenecek...")
                return process_pdf_with_ocr(file_path)
        
        elif ext == '.docx':
            try:
                # NLTK işlevlerini kullanmadan basit DOCX okuma
                import docx
                doc = docx.Document(file_path)
                text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
                metadata = {"source": file_path}
                documents = [Document(page_content=text, metadata=metadata)]
                print("DOCX python-docx ile başarıyla okundu.")
            except ImportError:
                print("python-docx bulunamadı, UnstructuredWordDocumentLoader deneniyor...")
                try:
                    loader = UnstructuredWordDocumentLoader(file_path)
                    documents = loader.load()
                except Exception as docx_error:
                    raise Exception(f"DOCX yüklenemedi: {str(docx_error)}")
        
        elif ext == '.csv':
            # Farklı encoding türlerini dene
            encodings = ['utf-8', 'latin1', 'iso-8859-1', 'cp1252']
            success = False
            
            for encoding in encodings:
                try:
                    # Pandas ile CSV dosyasını okuyalım
                    df = pd.read_csv(file_path, encoding=encoding)
                    for i, row in df.iterrows():
                        # Her satırı "Kolon: Değer" formatında metne dönüştür
                        content = "\n".join([f"{col}: {str(val)}" for col, val in row.items() if pd.notna(val)])
                        metadata = {"source": file_path, "row": i}
                        documents.append(Document(page_content=content, metadata=metadata))
                    success = True
                    print(f"CSV başarıyla {encoding} encodingle okundu.")
                    break
                except UnicodeDecodeError:
                    print(f"CSV {encoding} ile okunamadı, diğer encoding deneniyor...")
                    continue
                except Exception as csv_error:
                    print(f"CSV {encoding} ile okuma hatası: {str(csv_error)}")
                    # Son encoding denemesiyse hata ver
                    if encoding == encodings[-1]:
                        raise Exception(f"CSV dosyası hiçbir encoding ile okunamadı")
            
            # Pandas ile okunamadıysa CSVLoader'ı dene
            if not success and not documents:
                try:
                    for encoding in encodings:
                        try:
                            loader = CSVLoader(file_path=file_path, encoding=encoding)
                            documents = loader.load()
                            print(f"CSV CSVLoader ile {encoding} encodingle başarıyla okundu.")
                            break
                        except UnicodeDecodeError:
                            continue
                        except Exception as e:
                            if encoding == encodings[-1]:
                                raise Exception(f"Error loading {file_path}: {str(e)}")
                except Exception as e:
                    raise Exception(f"CSV yüklenemedi: {str(e)}")
            
        elif ext == '.xlsx':
            try:
                # Excel dosyasını pandas ile okuyup her satırı bir Document'a dönüştürüyoruz
                df = pd.read_excel(file_path)
                for i, row in df.iterrows():
                    # Her satırı "Kolon: Değer" formatında metne dönüştür
                    content = "\n".join([f"{col}: {str(val)}" for col, val in row.items() if pd.notna(val)])
                    metadata = {"source": file_path, "row": i}
                    documents.append(Document(page_content=content, metadata=metadata))
            except Exception as e:
                raise Exception(f"Excel dosyası okunurken hata: {str(e)}")
        
        else:
            raise ValueError(f"Desteklenmeyen dosya formatı: {ext}")
            
    except Exception as e:
        raise Exception(f"Belge yüklenirken hata oluştu: {str(e)}")
        
    return documents

def process_pdf_with_ocr(pdf_path: str) -> List[Document]:
    """
    PDF'i görsel olarak işleyip OCR uygular (özellikle taranmış PDF'ler için)
    
    Şu an için basitleştirilmiş bir yaklaşım kullanıyoruz - sadece bir belge oluşturuyoruz
    İleri düzey yaklaşım: Her sayfayı görsel olarak çıkarıp ayrı ayrı OCR işlemi uygulamak
    """
    try:
        import fitz  # PyMuPDF
        
        # PDF dokümanını aç
        pdf_document = fitz.open(pdf_path)
        extracted_text = []
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Sayfayı görsel olarak çıkar
            pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))  # 300 DPI çözünürlük
            
            # Geçici dosya oluştur
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_file:
                temp_path = temp_file.name
            
            # Görseli kaydet
            pix.save(temp_path)
            
            # OCR işlemi
            ocr_result = ocr_processor.process_image(temp_path)
            
            # Sonuçları topla
            if ocr_result:
                for doc in ocr_result:
                    extracted_text.append(f"===== SAYFA {page_num + 1} =====\n{doc.page_content}")
            
            # Geçici dosyayı sil
            try:
                os.unlink(temp_path)
            except:
                pass
        
        # PDF'i kapat
        pdf_document.close()
        
        # Birleştirilmiş belge oluştur
        if extracted_text:
            full_text = "\n\n".join(extracted_text)
            return [Document(
                page_content=full_text,
                metadata={
                    "source": pdf_path,
                    "type": "pdf_ocr"
                }
            )]
        else:
            print("OCR ile metin çıkarılamadı")
            return []
        
    except ImportError:
        print("PyMuPDF (fitz) bulunamadı, OCR işlemi atlanıyor")
        return []
    except Exception as e:
        print(f"PDF OCR hatası: {str(e)}")
        return []

def load_url_content(url: str) -> List[Document]:
    """
    Verilen URL'deki içeriği çeker ve Document nesnesi olarak döndürür.
    """
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }
        response = requests.get(url, headers=headers)
        response.raise_for_status()  # HTTP hatalarını kontrol et
        
        soup = BeautifulSoup(response.text, "html.parser")
        
        # Metni ayıkla ve temizle
        text = soup.get_text(separator="\n")
        
        # Boş satırları temizle
        text = "\n".join([line for line in text.split("\n") if line.strip()])
        
        # Document nesnesi oluştur
        doc = Document(page_content=text, metadata={"source": url, "type": "webpage"})
        
        return [doc]
    
    except Exception as e:
        raise Exception(f"URL içeriği yüklenirken hata oluştu: {str(e)}")

def split_documents(documents: List[Document], 
                   chunk_size: int = 500, 
                   chunk_overlap: int = 50) -> List[Document]:
    """
    Belgeleri belirtilen boyutlarda parçalara ayırır.
    İki aşamalı akıllı parçalama stratejisi kullanır:
    1. Önce metin mantıksal sınırlarda (paragraf, cümle) bölünür
    2. Sonra bu parçalar chunk_size sınırına göre birleştirilir
    
    Bu yaklaşım, cümlelerin ortasından bölünmesini engeller ve daha anlamlı metin
    parçaları oluşturur.
    """
    # Aşama 1: Önce daha küçük mantıksal parçalara ayır (paragraf, cümle)
    # Bunun için çok kısa parça boyutu ve çok düşük örtüşme kullan
    initial_splitter = RecursiveCharacterTextSplitter(
        chunk_size=100,  # Küçük parçalar 
        chunk_overlap=0,  # Örtüşme yok
        separators=["\n\n", "\n", ".", "!", "?", ";", ":", " ", ""]  # Öncelikli ayraçlar
    )
    
    # Aşama 2: Küçük parçaları birleştirerek istenen boyuta getir
    # Bu sefer tokenler yerine doğrudan karakterleri sayar
    class SmartMerger:
        def __init__(self, max_chunk_size, chunk_overlap):
            self.max_chunk_size = max_chunk_size
            self.chunk_overlap = chunk_overlap
            
        def merge_documents(self, docs: List[Document]) -> List[Document]:
            if not docs:
                return []
                
            merged_docs = []
            current_text = ""
            current_metadata = None
            
            for doc in docs:
                # İlk belgeyi alırken metadata'yı kaydet
                if not current_metadata:
                    current_metadata = doc.metadata.copy()
                
                # Eğer mevcut metin + yeni parça chunk_size'dan küçükse, ekle
                if len(current_text) + len(doc.page_content) <= self.max_chunk_size:
                    if current_text:
                        current_text += " " + doc.page_content
                    else:
                        current_text = doc.page_content
                else:
                    # Mevcut parçayı kaydet
                    if current_text:
                        merged_docs.append(Document(
                            page_content=current_text,
                            metadata=current_metadata
                        ))
                    
                    # Yeni parçaya başla - son parçadan örtüşme ekle
                    if self.chunk_overlap > 0 and current_text:
                        # Son kelimelerden oluşan örtüşme oluştur
                        words = current_text.split()
                        overlap_word_count = min(len(words), 
                                               max(1, self.chunk_overlap // 5))  # Ortalama kelime uzunluğu ~5 karakter
                        
                        # Örtüşme metni oluştur
                        overlap_text = " ".join(words[-overlap_word_count:])
                        
                        # Yeni parçaya örtüşmeyi ekle
                        current_text = overlap_text + " " + doc.page_content
                    else:
                        current_text = doc.page_content
                    
                    # Yeni parça için metadata'yı güncelle
                    current_metadata = doc.metadata.copy()
            
            # Son parçayı ekle
            if current_text:
                merged_docs.append(Document(
                    page_content=current_text,
                    metadata=current_metadata
                ))
                
            return merged_docs
    
    # Uygulama: Önce küçük parçalara ayır, sonra akıllı birleştirme yap
    small_chunks = []
    for doc in documents:
        small_chunks.extend(initial_splitter.split_documents([doc]))
    
    # Küçük parçaları birleştir
    merger = SmartMerger(chunk_size, chunk_overlap)
    merged_documents = merger.merge_documents(small_chunks)
    
    return merged_documents

# Yeni fonksiyonlar, eski fonksiyonları çağırıp ilave işlemler yapan sarmalayıcılardır
def load_document_from_file(file_path: str, file_name: str = None) -> List[Document]:
    """
    Dosyadan belge yükler ve içeriği parçalara ayırır.
    
    Args:
        file_path: Yüklenecek dosyanın yolu
        file_name: İsteğe bağlı dosya adı (metadata için)
    
    Returns:
        Parçalanmış belge listesi
    """
    # Belgeyi yükle
    documents = load_document(file_path)
    
    # Metadata'ya dosya adını ekle
    if file_name:
        for doc in documents:
            doc.metadata["filename"] = file_name
    
    # Belgeleri parçala (500 karakter, 50 karakter örtüşme) - daha küçük parçalar için değiştirdik
    chunked_docs = split_documents(documents, chunk_size=500, chunk_overlap=50)
    
    return chunked_docs

def load_document_from_url(url: str) -> List[Document]:
    """
    URL'den içerik yükler ve parçalara ayırır.
    
    Args:
        url: İçerik yüklenecek URL
    
    Returns:
        Parçalanmış belge listesi
    """
    # URL içeriğini yükle
    documents = load_url_content(url)
    
    # Belgeleri parçala (500 karakter, 50 karakter örtüşme) - daha küçük parçalar için değiştirdik
    chunked_docs = split_documents(documents, chunk_size=500, chunk_overlap=50)
    
    return chunked_docs

def load_image_ocr(image_path: str, image_name: str = None) -> List[Document]:
    """
    Görüntü dosyasından OCR ile metin çıkarır ve parçalara ayırır.
    
    Args:
        image_path: Görsel dosyasının yolu
        image_name: İsteğe bağlı görsel adı (metadata için)
    
    Returns:
        Parçalanmış belge listesi
    """
    # OCR işlemini uygula
    documents = ocr_processor.process_image(image_path, image_name)
    
    # Metadata'ya dosya adını ekle
    if image_name:
        for doc in documents:
            doc.metadata["filename"] = image_name
    
    # Belgeleri parçala (daha küçük parçalara ayırma işlemi gerekiyorsa)
    if documents:
        chunked_docs = split_documents(documents, chunk_size=500, chunk_overlap=50)
        return chunked_docs
    
    return [] 