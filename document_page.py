import os
import streamlit as st
import pandas as pd
import docx
from PyPDF2 import PdfReader
import pickle
from streamlit_pdf_viewer import pdf_viewer
from dotenv import load_dotenv
from langchain.chat_models import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain

def load_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

# CSS dosyasını yükle
load_css("style.css")

load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

if api_key is None:
    raise ValueError("OPENAI_API_KEY çevre değişkeni bulunamadı. Lütfen .env dosyasını kontrol edin.")

os.environ["OPENAI_API_KEY"] = api_key

# Vektör tabanı ve chat modelini bir kez yükle
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=20,
    length_function=len
)
embeddings = OpenAIEmbeddings()
llm = ChatOpenAI(temperature=0.07, model_name="gpt-3.5-turbo")
chain = load_qa_chain(llm=llm, chain_type="stuff")

def document_page():
    column1, column2 = st.columns(2)

    consultech_logo = "./images/consultech.png"
    st.sidebar.image(consultech_logo, width=200)
    uploaded_files = st.sidebar.file_uploader("Dosya Yükleyin", type=['pdf', 'csv', 'xlsx', 'xls', 'docx'],
                                              accept_multiple_files=True)

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    if "show_history" not in st.session_state:
        st.session_state.show_history = False

    if "current_response" not in st.session_state:
        st.session_state.current_response = None

    with column1:
        st.header("💬 Dokümanlarla Chat")
        if uploaded_files:
            # Yüklenen dosyaların işlenmesi ve vektör tabanının oluşturulması
            text = ""
            for file in uploaded_files:
                extension = file.name.split('.')[-1]
                if extension == "pdf":
                    file_reader = PdfReader(file)
                    for page in file_reader.pages:
                        text += page.extract_text()
                elif extension == "csv":
                    file_reader = pd.read_csv(file)
                    text += "\n".join(file_reader.apply(lambda row: ', '.join(row.values.astype(str)), axis=1))
                elif extension in ["xlsx", "xls"]:
                    file_reader = pd.read_excel(file)
                    text += "\n".join(file_reader.apply(lambda row: ', '.join(row.values.astype(str)), axis=1))
                elif extension == "docx":
                    file_reader = docx.Document(file)
                    text += ' '.join([paragraph.text for paragraph in file_reader.paragraphs])

            if text:
                chunks = text_splitter.split_text(text=text)
                vector_store = FAISS.from_texts(chunks, embedding=embeddings)

                # Vektör tabanını kaydet
                vector_base_name = "vector_base.pkl"
                with open(vector_base_name, "wb") as f:
                    pickle.dump(vector_store, f)
            else:
                st.error("Bir şeyler yanlış gitti")
        else:
            st.write("Merhaba! Doküman yüklemediniz. Lütfen doküman yükleyin.")

        # Doküman olmasa bile soruları sorabileceğimiz alanı ve butonları göster
        query = st.text_input("Dokümanlara Ne Sormak İstiyorsunuz?", label_visibility="visible", key="query_input")
        send_button = st.button("Gönder")

        if send_button:
            if query and os.path.exists("vector_base.pkl"):
                with open("vector_base.pkl", "rb") as f:
                    vector_store = pickle.load(f)

                    docs = vector_store.similarity_search(query=query, k=10)
                    response = chain.run(input_documents=docs, question=query)

                    # Sohbet geçmişine ekle
                    st.session_state.chat_history.append({
                        "soru": query,
                        "cevap": response
                    })

                    # Cevabı kaydet
                    st.session_state.current_response = response

        # Cevabı görüntüle (varsayılan olarak boş olsa bile yerini koruyarak)
        st.divider()
        st.subheader("Cevap:")
        if st.session_state.current_response:
            st.write(st.session_state.current_response)
        else:
            st.write("Henüz bir cevap yok.")
        st.divider()

        # Sohbet geçmişini görüntüleme butonu
        if st.button("📜 Sohbet Geçmişini Göster/Gizle"):
            st.session_state.show_history = not st.session_state.show_history

        # Sohbet geçmişini görüntüle
        if st.session_state.show_history and st.session_state.chat_history:
            st.subheader("Sohbet Geçmişi:")
            for idx, item in enumerate(st.session_state.chat_history, 1):
                st.write(f"**{idx}.**")
                st.write(f"**Siz:** {item['soru']}")
                st.write(f"**Asistan:** {item['cevap']}")

    with column2:
        st.subheader("Yüklenen Dokümanlar:")
        if uploaded_files:
            for file in uploaded_files:
                st.write(file.name)
                file_extension = file.name.split(".")[-1]
                if file_extension == "pdf":
                    pdf_reader = PdfReader(file)
                    binary_data = file.getvalue()
                    pdf_viewer(input=binary_data, width=500, height=550)
                elif file_extension == "xlsx":
                    df = pd.read_excel(file)
                    st.write(df)
                elif file_extension == "csv":
                    try:
                        df = pd.read_csv(file)
                        st.write(df)
                    except pd.errors.EmptyDataError:
                        st.write("Dosya Yükleme Başarılı.")
                elif file_extension == "docx":
                    docx_reader = docx.Document(file)
                    for paragraph in docx_reader.paragraphs:
                        st.write(paragraph.text)
                else:
                    st.write("Dosya Türü Desteklenmiyor!")
        else:
            st.write("Henüz doküman yüklemediniz.")
