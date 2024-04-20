import os
import streamlit as st
import pandas as pd
import docx
from PyPDF2 import PdfReader
import pickle
from langchain.chat_models import ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from streamlit_pdf_viewer import pdf_viewer


os.environ["OPENAI_API_KEY"] = "sk-proj-ugbMH2Kt5r4gSJCGsHFQT3BlbkFJhIaQPQ7uIVBmx7fdR96H"
openai_api_key = os.getenv("OPENAI_API_KEY")
if openai_api_key:
    os.environ["OPENAI_API_KEY"] = openai_api_key
else:
    st.error("OpenAI API key is missing. Please make sure to set it in your .env file.")

# Vektör tabanı ve chat modelini bir kez yükle
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=20,
    length_function=len
)

embeddings = OpenAIEmbeddings()
llm = ChatOpenAI(temperature=0.07, model_name="gpt-3.5-turbo")
chain = load_qa_chain(llm=llm, chain_type="stuff")

# Sayfayı ikiye bölmek için 2 sütun oluştur
column1, column2 = st.columns(2)

# Sol sütun üzerinde işlemler
import streamlit as st

def pdf_page():
    st.write("PDF içerikleri burada gösterilir...")
    # Diğer PDF ile ilgili işlemler


# Sayfayı ikiye bölme
column1, column2 = st.columns(2)

st.markdown("""
    <style>
        .st-emotion-cache-gh2jqd {
            max-width: none !important;
            padding: 2rem !important;
        }
    </style>
    """, unsafe_allow_html=True)

image_file = "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjcuMTE1IiBoZWlnaHQ9IjE5LjY4NyIgdmlld0JveD0iMCAwIDEyNy4xMTUgMTkuNjg3Ij4KICA8ZyBpZD0iR3JvdXBfMjE2MCIgZGF0YS1uYW1lPSJHcm91cCAyMTYwIiB0cmFuc2Zvcm09InRyYW5zbGF0ZSgtNDY2LjM0NSAtMzc4LjU1MikiPgogICAgPGcgaWQ9Ikdyb3VwXzIxNTQiIGRhdGEtbmFtZT0iR3JvdXAgMjE1NCIgdHJhbnNmb3JtPSJ0cmFuc2xhdGUoNDY2LjM0NSAzNzguNTUyKSI+CiAgICAgIDxwYXRoIGlkPSJQYXRoXzc4OSIgZGF0YS1uYW1lPSJQYXRoIDc4OSIgZD0iTTQ4My40NjIsMzc4LjU1MWwtNi43ODksMTkuNjg2aC00LjE4MWwtNi4xNDYtMTkuNjg2aDQuODI0bDQuOTY3LDE2LjAyNyw1LjQ2Ny0xNi4wMjdoMS44NTgiIHRyYW5zZm9ybT0idHJhbnNsYXRlKC00NjYuMzQ1IC0zNzguNTUxKSIgZmlsbD0iIzIwMmUzNiIvPgogICAgICA8cGF0aCBpZD0iUGF0aF83OTAiIGRhdGEtbmFtZT0iUGF0aCA3OTAiIGQ9Ik01MzYuNjk1LDM5Mi4xaDYuOWwtMy4zNTktOS45NzYtMy41MzgsOS45NzZtNy4wMzktMTMuNTQ3LDYuNzUzLDE5LjY4NmgtNC44NTlsLTEuNTczLTQuNjkzaC03Ljg2MWwtMS42NDQsNC42OTNoLTEuOTI5bDcuMDM5LTE5LjY4NloiIHRyYW5zZm9ybT0idHJhbnNsYXRlKC01MTMuMjE5IC0zNzguNTUyKSIgZmlsbD0iIzIwMmUzNiIvPgogICAgICA8cGF0aCBpZD0iUGF0aF83OTEiIGRhdGEtbmFtZT0iUGF0aCA3OTEiIGQ9Ik02MjkuMjM1LDM3OC41NTF2MTkuNjg2aC0zLjY4bC0xMC40MzUtMTUuMDIzdjE1LjAyM2gtMS40NjRWMzc4LjU1MWgzLjYxbDEwLjUwNiwxNS4wMjNWMzc4LjU1MWgxLjQ2NCIgdHJhbnNmb3JtPSJ0cmFuc2xhdGUoLTU3MC41MyAtMzc4LjU1MSkiIGZpbGw9IiMyMDJlMzYiLz4KICAgICAgPHBhdGggaWQ9IlBhdGhfNzkyIiBkYXRhLW5hbWU9IlBhdGggNzkyIiBkPSJNNzA1LjIyOSwzNzguNTUxdjEuNTA1SDY5Ni43djYuNjExaDcuMDA2djEuNDQ2SDY5Ni43djguNjc4aDguNTI0djEuNDQ2SDY5MS44OFYzNzguNTUxaDEzLjM0OCIgdHJhbnNmb3JtPSJ0cmFuc2xhdGUoLTYyNS44NTYgLTM3OC41NTEpIiBmaWxsPSIjMjAyZTM2Ii8+CiAgICAgIDxwYXRoIGlkPSJQYXRoXzc5MyIgZGF0YS1uYW1lPSJQYXRoIDc5MyIgZD0iTTc2My4zOCwzNzguNTUxdjE4LjIxMWg4LjY1NHYxLjQ3NUg3NTguNjYzVjM3OC41NTFoNC43MTciIHRyYW5zZm9ybT0idHJhbnNsYXRlKC02NzMuMDg4IC0zNzguNTUxKSIgZmlsbD0iIzIwMmUzNiIvPgogICAgICA8cGF0aCBpZD0iUGF0aF83OTQiIGRhdGEtbmFtZT0iUGF0aCA3OTQiIGQ9Ik04MjUuNywzNzguNTUxdjE4LjIxMWg4LjY1NHYxLjQ3NUg4MjAuOTgxVjM3OC41NTFIODI1LjciIHRyYW5zZm9ybT0idHJhbnNsYXRlKC03MTcuMTYyIC0zNzguNTUxKSIgZmlsbD0iIzIwMmUzNiIvPgogICAgPC9nPgogICAgPGcgaWQ9Ikdyb3VwXzIxNTUiIGRhdGEtbmFtZT0iR3JvdXAgMjE1NSIgdHJhbnNmb3JtPSJ0cmFuc2xhdGUoNTg4Ljc3NCAzNzguNTUzKSI+CiAgICAgIDxyZWN0IGlkPSJSZWN0YW5nbGVfODQ3IiBkYXRhLW5hbWU9IlJlY3RhbmdsZSA4NDciIHdpZHRoPSI0LjY4NyIgaGVpZ2h0PSIxLjQ2MiIgdHJhbnNmb3JtPSJ0cmFuc2xhdGUoMCAwKSIgZmlsbD0iIzIwMmUzNiIvPgogICAgICA8cmVjdCBpZD0iUmVjdGFuZ2xlXzg0OCIgZGF0YS1uYW1lPSJSZWN0YW5nbGUgODQ4IiB3aWR0aD0iNC42ODciIGhlaWdodD0iMTYuOTAzIiB0cmFuc2Zvcm09InRyYW5zbGF0ZSgwIDIuNzMyKSIgZmlsbD0iIzIwMmUzNiIvPgogICAgPC9nPgogIDwvZz4KPC9zdmc+Cg=="
st.sidebar.image(image_file, use_column_width=True)
uploadedFiles = st.sidebar.file_uploader("",type=['pdf', '.csv', '.xlsx', '.xls', '.docx'], accept_multiple_files=True)

# Sol sütunda pdf_page() fonksiyonunu çalıştırma
with column1:
    st.header("PDF Dokümanı")
# Yüklenen dosyaların işlenmesi ve vektör tabanının oluşturulması
    if uploadedFiles:
        text = ""
        for file in uploadedFiles:
            extension = file.name[len(file.name) - 3:]
            if (extension == "pdf"):
                file_reader = PdfReader(file)
                for page in file_reader.pages:
                    text += page.extract_text()
            elif (extension == "csv"):
                file_reader = pd.read_csv(file)
                text += "\n".join(
                    file_reader.apply(lambda row: ', '.join(row.values.astype(str)), axis=1))
            elif (extension == "lsx" or extension == "xls"):
                file_reader = pd.read_excel(file)
                text += "\n".join(
                    file_reader.apply(lambda row: ', '.join(row.values.astype(str)), axis=1))
            elif (extension == "ocx"):
                file_reader = docx.Document(file)
                list = [paragraph.text for paragraph in file_reader.paragraphs]
                text += ' '.join(list)

        if (text):
            chunks = text_splitter.split_text(text=text)
            VectorStore = FAISS.from_texts(chunks, embedding=embeddings)

            # Vektör tabanının kaydedilmesi
            vector_base_name = "vector_base.pkl"
            with open(vector_base_name, "wb") as f:
                pickle.dump(VectorStore, f)

        else:
            st.error("Bir şeyler yanlış gitti")

    # Chat modelini yükle ve kullanıcı ile etkileşime geç
    if uploadedFiles:
        st.write("Dokümanlara Ne Sormak İstiyorsunuz?")
        query = st.text_input("Soruyu aşağıya yazınız ve enter tuşuna basımız.")

        # Enter tuşuna basıldığında
        if query and os.path.exists(vector_base_name):
            with open(vector_base_name, "rb") as f:
                VectorStore = pickle.load(f)

                # Benzer dokümanların getirilmesi
                k = 10  # Number of nearest neighbors to retrieve
                distances = []  # List to store the distances
                labels = []
                docs = VectorStore.similarity_search(
                    query=query, k=k, distances=distances, labels=labels)

                # Chat modelinin yüklenmesi ve dokümanla konuşulması
                response = chain.run(input_documents=docs, question=query)
                st.divider()
                st.subheader("Cevap: ")
                st.write(response)
                st.divider()

    else:
        st.write("Merhaba! Doküman yüklemediniz. Lütfen doküman yükleyin.")

# Sağ sütun için ayrı işlemler
with column2:
# Yuklenen belgelerin içeriği
    if uploadedFiles:
        st.subheader("Yüklenen Dokümanlar:")
        for file in uploadedFiles:
            st.write(file.name)
            file_extension = file.name.split(".")[-1]  # Dosya uzantısını al
            if file_extension == "pdf":
                pdf_reader = PdfReader(file)
                binary_data = file.getvalue()
                pdf_viewer(input=binary_data, width=600, height=500)
            elif file_extension == "xlsx":
                df = pd.read_excel(file)
                st.write(df)
            elif file_extension == "csv":
                try:
                    df = pd.read_csv(file)
                    st.write(df)
                except pd.errors.EmptyDataError:
                    st.write("Dosya Yükleme Başarılı.")
            elif file_extension == "docx":
                docx_reader = docx.Document(file)
                for paragraph in docx_reader.paragraphs:
                    st.write(paragraph.text)
            else:
                st.write("Dosya Türü Desteklenmiyor!")

